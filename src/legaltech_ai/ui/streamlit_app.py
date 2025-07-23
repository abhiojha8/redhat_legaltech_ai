"""
Simple Streamlit application for LegalTech AI.
"""

import streamlit as st
import os
from dotenv import load_dotenv
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(__file__)))
from watsonx_service import WatsonxService
from rag_service import RAGService
from call_data_service import CallDataService

# Load environment variables
load_dotenv()

def setup_page():
    """Setup Streamlit page configuration."""
    st.set_page_config(
        page_title="🧠 LegalTech AI",
        page_icon="⚖️",
        layout="wide"
    )
    
    st.title("🧠 LegalTech AI - Red Hat Hackathon")
    st.markdown("**Powered by IBM watsonx.ai**")
    st.markdown("---")

def document_upload_section():
    """Handle document upload and processing."""
    st.header("📄 Document Analysis")
    
    uploaded_file = st.file_uploader(
        "Upload a document (PDF, DOCX, TXT)",
        type=['pdf', 'docx', 'txt'],
        help="Maximum file size: 10MB"
    )
    
    if uploaded_file:
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        
        # Simple text extraction
        try:
            if uploaded_file.type == "text/plain":
                text = str(uploaded_file.read(), "utf-8")
            elif uploaded_file.type == "application/pdf":
                try:
                    import pypdf
                    pdf_reader = pypdf.PdfReader(uploaded_file)
                    text = ""
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    if not text.strip():
                        st.warning("⚠️ No text could be extracted from this PDF. It might be image-based.")
                except ImportError:
                    st.error("PDF processing requires pypdf package")
                    return None
                except Exception as pdf_error:
                    st.error(f"PDF processing error: {str(pdf_error)}")
                    return None
            elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                try:
                    import docx
                    doc = docx.Document(uploaded_file)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                except ImportError:
                    st.error("DOCX processing requires python-docx package")
                    return None
            else:
                st.error("Unsupported file type")
                return None
            
            st.success(f"✅ Document loaded ({len(text)} characters)")
            
            # Show preview
            with st.expander("📖 Document Preview"):
                preview = text[:500] + "..." if len(text) > 500 else text
                st.text(preview)
            
            return text
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return None
    
    return None

def chat_section():
    """Handle chat functionality with RAG-based document context."""
    st.header("💬 Legal Assistant Chat with RAG")
    
    # Initialize RAG service
    if "rag_service" not in st.session_state:
        st.session_state.rag_service = RAGService()
    
    rag_service = st.session_state.rag_service
    
    # Document upload section for RAG
    with st.expander("📄 Upload Document for RAG", expanded=False):
        st.info("Upload a document to create a searchable knowledge base. Large documents are automatically chunked and vectorized.")
        
        context_file = st.file_uploader(
            "Upload document for RAG knowledge base",
            type=['pdf', 'docx', 'txt'],
            help="Document will be chunked and stored in vector database for efficient retrieval",
            key="rag_context_file"
        )
        
        if context_file:
            context_file.seek(0)
            with st.spinner("🔍 Processing document for RAG..."):
                try:
                    # Extract text
                    document_text = ""
                    if context_file.type == "text/plain":
                        document_text = str(context_file.read(), "utf-8")
                    elif context_file.type == "application/pdf":
                        import pypdf
                        pdf_reader = pypdf.PdfReader(context_file)
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                document_text += page_text + "\n"
                    elif context_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                        import docx
                        doc = docx.Document(context_file)
                        document_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    
                    if document_text.strip():
                        # Add to RAG system
                        result = rag_service.add_document(document_text, context_file.name)
                        
                        if result["success"]:
                            st.success(f"✅ Document processed! Created {result['chunks_created']} searchable chunks")
                            st.info(f"📊 Document: {result['document_name']} ({len(document_text)} characters)")
                        else:
                            st.error(f"❌ RAG processing failed: {result['error']}")
                    else:
                        st.warning("⚠️ No text could be extracted from this document")
                        
                except Exception as e:
                    st.error(f"Error processing document: {str(e)}")
    
    # Show RAG status
    if rag_service.has_documents():
        st.success("🔍 RAG Knowledge Base: Active")
        col1, col2 = st.columns([3, 1])
        with col1:
            st.info("Your questions will be answered using relevant chunks from the uploaded document")
        with col2:
            if st.button("🗑️ Clear Knowledge Base"):
                rag_service.clear_documents()
                st.rerun()
    else:
        st.warning("📚 No documents in knowledge base - upload a document above for context-aware answers")
    
    # Initialize chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask a legal question..."):
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.write(prompt)
        
        # Get AI response with RAG context
        try:
            service = WatsonxService()
            with st.chat_message("assistant"):
                with st.spinner("🔍 Searching knowledge base..."):
                    # Get relevant context using RAG
                    if rag_service.has_documents():
                        context = rag_service.get_context_for_query(prompt)
                        context_used = bool(context.strip())
                    else:
                        context = ""
                        context_used = False
                    
                    response = service.chat_response(prompt, context)
                    
                    if response['status'] == 'success':
                        answer = response['response']
                        st.write(answer)
                        st.session_state.messages.append({"role": "assistant", "content": answer})
                        
                        # Show context indicators
                        if context_used:
                            st.caption("🔍 Answer based on relevant document chunks via RAG")
                        else:
                            st.caption("💭 General knowledge answer (no document context)")
                    else:
                        error_msg = f"Error: {response['error']}"
                        st.error(error_msg)
                        st.session_state.messages.append({"role": "assistant", "content": error_msg})
        except Exception as e:
            error_msg = f"Service error: {str(e)}"
            st.error(error_msg)
            st.session_state.messages.append({"role": "assistant", "content": error_msg})

def call_data_analysis_section():
    """Handle call data analysis functionality."""
    st.header("📊 Call Data Penalty Analysis")
    st.markdown("Upload Excel call data files and analyze them for potential penalties using AI and regulatory context.")
    
    # Initialize services
    if "call_data_service" not in st.session_state:
        st.session_state.call_data_service = CallDataService()
    if "rag_service" not in st.session_state:
        st.session_state.rag_service = RAGService()
    
    call_service = st.session_state.call_data_service
    rag_service = st.session_state.rag_service
    
    # Regulatory context section
    with st.expander("📋 Upload Regulatory Documents (Optional)", expanded=False):
        st.info("Upload telecommunications regulations that will be used as context for penalty analysis")
        
        reg_file = st.file_uploader(
            "Upload regulatory document",
            type=['pdf', 'docx', 'txt'],
            help="This document will provide regulatory context for penalty analysis",
            key="regulatory_context"
        )
        
        if reg_file:
            reg_file.seek(0)
            with st.spinner("🔍 Processing regulatory document..."):
                try:
                    # Extract text
                    reg_text = ""
                    if reg_file.type == "text/plain":
                        reg_text = str(reg_file.read(), "utf-8")
                    elif reg_file.type == "application/pdf":
                        import pypdf
                        pdf_reader = pypdf.PdfReader(reg_file)
                        for page in pdf_reader.pages:
                            page_text = page.extract_text()
                            if page_text:
                                reg_text += page_text + "\n"
                    elif reg_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                        import docx
                        doc = docx.Document(reg_file)
                        reg_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    
                    if reg_text.strip():
                        result = rag_service.add_document(reg_text, reg_file.name)
                        if result["success"]:
                            st.success(f"✅ Regulatory document processed! Created {result['chunks_created']} searchable chunks")
                        else:
                            st.error(f"❌ Processing failed: {result['error']}")
                    
                except Exception as e:
                    st.error(f"Error processing regulatory document: {str(e)}")
    
    # Show regulatory context status
    if rag_service.has_documents():
        st.success("📖 Regulatory Context: Active")
        col1, col2 = st.columns([3, 1])
        with col1:
            st.info("Penalty analysis will reference uploaded regulations")
        with col2:
            if st.button("🗑️ Clear Regulations", key="clear_reg"):
                rag_service.clear_documents()
                st.rerun()
    else:
        st.warning("📚 No regulatory context loaded - analysis will use general knowledge")
    
    # Call data upload section
    st.markdown("---")
    st.subheader("📈 Upload Call Data")
    
    excel_file = st.file_uploader(
        "Upload Excel call data file",
        type=['xlsx', 'xls'],
        help="Upload Excel file containing call records for penalty analysis"
    )
    
    if excel_file:
        with st.spinner("📊 Processing call data..."):
            # Process Excel file
            result = call_service.process_excel_file(excel_file)
            
            if result["success"]:
                df = result["data"]
                stats = result["stats"]
                quality_issues = result["quality_issues"]
                
                st.success(f"✅ Call data processed successfully!")
                
                # Display statistics
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Total Records", stats["total_records"])
                with col2:
                    st.metric("Date Range", f"{stats['date_range']['duration_days']} days")
                with col3:
                    st.metric("Columns", len(stats["columns"]))
                
                # Show data preview
                with st.expander("📋 Data Preview", expanded=True):
                    st.dataframe(df.head(10), use_container_width=True)
                
                # Show data quality issues
                if quality_issues:
                    with st.expander("⚠️ Data Quality Issues", expanded=False):
                        for issue in quality_issues:
                            st.warning(f"• {issue}")
                
                # Analyze for violations
                if st.button("🔍 Analyze for Penalties", type="primary"):
                    with st.spinner("🤖 Analyzing call data for violations..."):
                        # Run compliance analysis
                        violation_result = call_service.analyze_compliance_violations(df)
                        
                        if violation_result["success"]:
                            violations = violation_result["violations"]
                            
                            # Display violation summary
                            st.markdown("### 📊 Penalty Risk Assessment")
                            
                            summary = violations["summary"]
                            col1, col2, col3 = st.columns(3)
                            
                            with col1:
                                st.metric("High Risk", summary["high_risk_count"], delta=-summary["high_risk_count"] if summary["high_risk_count"] > 0 else None)
                            with col2:
                                st.metric("Medium Risk", summary["medium_risk_count"], delta=-summary["medium_risk_count"] if summary["medium_risk_count"] > 0 else None)
                            with col3:
                                st.metric("Low Risk", summary["low_risk_count"])
                            
                            # Additional TRAI-specific metrics
                            if summary.get("estimated_penalty_inr", 0) > 0:
                                st.error(f"💰 **Estimated TRAI Penalty: ₹{summary['estimated_penalty_inr']:,}**")
                            
                            if summary.get("call_drop_violation_detected", False):
                                st.warning("📞 **Call Drop Rate Violation Detected** - TRAI's highest priority violation!")
                            
                            # Show TRAI context
                            with st.expander("📋 TRAI Penalty Context", expanded=False):
                                st.info(f"**Regulatory Authority**: {summary.get('regulatory_authority', 'TRAI')}")
                                st.info(f"**Standards Year**: {summary.get('benchmark_year', '2024')}")
                                
                                if 'key_trai_limits' in summary:
                                    limits = summary['key_trai_limits']
                                    st.write("**Key TRAI Limits:**")
                                    for key, value in limits.items():
                                        st.write(f"• {key.replace('_', ' ').title()}: {value}")
                            
                            # Display violations by risk level
                            for risk_level, color, icon in [
                                ("high_risk", "🔴", "🚨"),
                                ("medium_risk", "🟡", "⚠️"), 
                                ("low_risk", "🟢", "ℹ️")
                            ]:
                                if violations[risk_level]:
                                    st.markdown(f"#### {icon} {risk_level.replace('_', ' ').title()} Violations")
                                    
                                    for violation in violations[risk_level]:
                                        with st.container():
                                            st.markdown(f"**{violation['type']}**")
                                            st.write(violation['description'])
                                            
                                            # Show sample records for high risk
                                            if risk_level == "high_risk" and 'sample_records' in violation:
                                                with st.expander("View Sample Records"):
                                                    st.json(violation['sample_records'])
                            
                            # AI Analysis with Regulatory Context
                            st.markdown("### 🤖 AI Penalty Analysis")
                            
                            with st.spinner("🔍 Getting regulatory context and AI analysis..."):
                                # Generate analysis prompt
                                context = ""
                                if rag_service.has_documents():
                                    context = rag_service.get_context_for_query("telecommunications penalties violations fines compliance")
                                
                                analysis_prompt = call_service.generate_analysis_prompt(violations, context)
                                
                                # Get AI analysis
                                watsonx_service = WatsonxService()
                                ai_response = watsonx_service.chat_response(
                                    "Analyze these telecommunications call data violations for potential penalties:",
                                    analysis_prompt
                                )
                                
                                if ai_response['status'] == 'success':
                                    st.markdown("#### 📝 Legal Analysis & Recommendations")
                                    st.write(ai_response['response'])
                                    
                                    if context:
                                        st.caption("🔍 Analysis includes regulatory context from uploaded documents")
                                    else:
                                        st.caption("💭 Analysis based on general telecommunications law knowledge")
                                else:
                                    st.error(f"AI analysis failed: {ai_response['error']}")
                            
                        else:
                            st.error(f"❌ Violation analysis failed: {violation_result['error']}")
                            
            else:
                st.error(f"❌ Excel processing failed: {result['error']}")

def main():
    """Main application."""
    setup_page()
    
    # Check environment variables
    required_vars = ['WATSONX_API_KEY', 'WATSONX_URL', 'WATSONX_PROJECT_ID']
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    
    if missing_vars:
        st.error(f"❌ Missing environment variables: {', '.join(missing_vars)}")
        st.info("Please check your .env file configuration")
        st.stop()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Configuration")
        model_id = os.getenv('MODEL_ID', 'meta-llama/llama-3-3-70b-instruct')
        st.info(f"Model: {model_id.split('/')[-1]}")
        st.markdown("---")
        
        mode = st.selectbox(
            "Choose Mode",
            ["Document Analysis", "Chat Assistant", "Call Data Analysis"]
        )
    
    # Main content based on mode
    if mode == "Document Analysis":
        document_text = document_upload_section()
        
        if document_text and st.button("🚀 Analyze Document", type="primary"):
            try:
                service = WatsonxService()
                with st.spinner("Analyzing document with watsonx.ai..."):
                    result = service.analyze_document(document_text)
                
                if result['status'] == 'success':
                    st.success("✅ Analysis completed!")
                    st.markdown("### 📊 Analysis Results")
                    st.text_area("Analysis", value=result['analysis'], height=400)
                else:
                    st.error(f"❌ Analysis failed: {result['error']}")
            except Exception as e:
                st.error(f"Error: {str(e)}")
    
    elif mode == "Chat Assistant":
        chat_section()
    
    elif mode == "Call Data Analysis":
        call_data_analysis_section()

if __name__ == "__main__":
    main()